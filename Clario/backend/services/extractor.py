"""
extractor.py
────────────
Responsible for pulling clean plain-text out of uploaded files.

Supports:
  • .pdf   – pdfplumber  (text-layer PDFs; raises clearly for scanned/image PDFs)
  • .docx  – python-docx (paragraphs + table cells)
  • .txt   – plain UTF-8 passthrough
  • raw str – direct sanitisation (pasted text path)

All public functions return a clean str or raise FastAPI HTTPException.
"""

from __future__ import annotations

import io
import logging
import re
import unicodedata
from pathlib import Path

import pdfplumber
from docx import Document
from fastapi import HTTPException

log = logging.getLogger(__name__)

# ── Constants ──────────────────────────────────────────────────────────────────
MAX_FILE_BYTES: int = 10 * 1024 * 1024          # 10 MB hard limit (PRD §5.1)
MIN_EXTRACTABLE_CHARS: int = 50                  # reject near-empty docs
ALLOWED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx", ".txt"})


# ══════════════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════════════

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch to the correct extractor based on file extension.

    Args:
        file_bytes: Raw bytes from the uploaded file.
        filename:   Original filename (used only for extension detection).

    Returns:
        Cleaned plain-text string ready for the AI service.

    Raises:
        HTTPException 400 – file too large.
        HTTPException 415 – unsupported extension.
        HTTPException 422 – extraction succeeded but yielded no usable text.
    """
    _enforce_size(file_bytes, filename)

    ext = Path(filename).suffix.lower()

    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=(
                f"Unsupported file type '{ext}'. "
                f"Accepted: {', '.join(sorted(ALLOWED_EXTENSIONS))}."
            ),
        )

    if ext == ".pdf":
        raw = _extract_pdf(file_bytes)
    elif ext == ".docx":
        raw = _extract_docx(file_bytes)
    else:                           # .txt
        raw = _extract_txt(file_bytes)

    return _post_process(raw, source=filename)


def extract_text_from_string(text: str) -> str:
    """
    Sanitise and return raw pasted text.

    Args:
        text: User-pasted document content.

    Returns:
        Cleaned plain-text string.

    Raises:
        HTTPException 422 – text is blank or too short after cleaning.
    """
    cleaned = _post_process(text, source="pasted text")
    return cleaned


# ══════════════════════════════════════════════════════════════════════════════
# Private – format-specific extractors
# ══════════════════════════════════════════════════════════════════════════════

def _extract_pdf(file_bytes: bytes) -> str:
    """Extract text from every page using pdfplumber."""
    pages: list[str] = []

    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if not pdf.pages:
                raise HTTPException(
                    status_code=422, detail="PDF has no pages."
                )
            for page_num, page in enumerate(pdf.pages, start=1):
                try:
                    page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                except Exception as page_err:
                    log.warning("PDF page %d extraction error: %s", page_num, page_err)
                    continue
                if page_text and page_text.strip():
                    pages.append(page_text.strip())

    except HTTPException:
        raise
    except Exception as exc:
        log.error("pdfplumber failed: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=422,
            detail=f"Failed to parse PDF: {exc}",
        )

    if not pages:
        raise HTTPException(
            status_code=422,
            detail=(
                "No extractable text found in this PDF. "
                "It may be a scanned / image-only document. "
                "Please copy-paste the text directly instead."
            ),
        )

    return "\n\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    """
    Extract text from paragraphs AND table cells of a DOCX file.
    Tables in legal contracts often contain important clause data.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        log.error("python-docx failed: %s", exc, exc_info=True)
        raise HTTPException(
            status_code=422,
            detail=f"Failed to parse DOCX: {exc}",
        )

    chunks: list[str] = []

    # Body paragraphs
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            chunks.append(stripped)

    # Table cells (preserve row structure)
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                chunks.append(" | ".join(row_cells))

    if not chunks:
        raise HTTPException(
            status_code=422,
            detail="DOCX appears to be empty or contains only images / embedded objects.",
        )

    return "\n\n".join(chunks)


def _extract_txt(file_bytes: bytes) -> str:
    """Decode a plain-text file (UTF-8 with BOM fallback)."""
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(
        status_code=422,
        detail="Could not decode text file. Ensure it is UTF-8 encoded.",
    )


# ══════════════════════════════════════════════════════════════════════════════
# Private – shared helpers
# ══════════════════════════════════════════════════════════════════════════════

def _enforce_size(file_bytes: bytes, filename: str) -> None:
    size_mb = len(file_bytes) / (1024 * 1024)
    if len(file_bytes) > MAX_FILE_BYTES:
        raise HTTPException(
            status_code=400,
            detail=(
                f"File '{filename}' is {size_mb:.1f} MB. "
                f"Maximum allowed size is {MAX_FILE_BYTES // (1024*1024)} MB."
            ),
        )


def _post_process(text: str, *, source: str) -> str:
    """
    Normalise whitespace, strip non-printable characters, and validate length.
    """
    # Normalise Unicode to NFC (handles ligatures, combining chars, etc.)
    text = unicodedata.normalize("NFC", text)

    # Remove non-printable control characters (except newlines/tabs)
    text = re.sub(r"[^\S\n\t ]+", " ", text)           # collapse odd whitespace
    text = re.sub(r"[ \t]{2,}", " ", text)              # collapse inline spaces
    text = re.sub(r"\n{3,}", "\n\n", text)              # max 2 consecutive newlines
    text = text.strip()

    if len(text) < MIN_EXTRACTABLE_CHARS:
        raise HTTPException(
            status_code=422,
            detail=(
                f"Extracted text from {source} is too short "
                f"({len(text)} chars). The document may be empty or corrupted."
            ),
        )

    log.info("Extracted %d chars from %s", len(text), source)
    return text
